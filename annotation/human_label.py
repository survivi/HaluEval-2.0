# coding: utf-8
import os
import json
from docx import Document
from docx.shared import Inches
from sample import check_exist
from docx.oxml.ns import qn
from docx.shared import Pt
from evaluate import read_json


def set_font(run):
    """
    Set font for Chinese characters.
    """
    run.font.name = "仿宋"
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")


def split_list(lst, x=4):
    """
    Divide a list evenly into x parts.
    """
    n = len(lst)
    k = n // x
    m = n % x
    return [lst[i * k + min(i, m) : (i + 1) * k + min(i + 1, m)] for i in range(x)]


def generate_doc(model, file, samples, save_path):
    """
    Generate a docx file for annotation.
    """
    document = Document()
    run = document.add_heading("", 0).add_run("幻象标注")
    run.font.size = Pt(20)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for i in range(len(samples)):
        table = document.add_table(rows=8, cols=2, style="Table Grid")
        col_width_dic = {0: 3, 1: 7}
        for col_num in range(2):
            table.cell(0, col_num).width = Inches(col_width_dic[col_num])

        hdr_cells = table.columns[0].cells
        hdr_cells[0].text = "ID"
        run = hdr_cells[1].paragraphs[0].add_run("领域")
        set_font(run)
        run = hdr_cells[2].paragraphs[0].add_run("用户问题")
        set_font(run)
        run = (
            hdr_cells[3]
            .paragraphs[0]
            .add_run("问题打分（逗号分隔）：\n可读性（1-5）\n规范性（1-5）\n具体性（1-5）")
        )
        set_font(run)
        run = hdr_cells[4].paragraphs[0].add_run("模型回复")
        set_font(run)
        run = hdr_cells[5].paragraphs[0].add_run("回复标注（二选一）：\n1-回复与问题相关\n2-回复与问题不相关")
        set_font(run)
        run = hdr_cells[6].paragraphs[0].add_run("提取的事实")
        set_font(run)
        run = (
            hdr_cells[7]
            .paragraphs[0]
            .add_run(
                "事实标注\n（每条事实八选一）：\n1-完全正确事实\n2-实体错误事实\n3-关系错误事实\n4-信息缺失事实\n5-信息过时事实\n6-表述绝对事实\n7-无法验证事实\n8-非事实表述"
            )
        )
        set_font(run)

        hdr_cells = table.columns[1].cells
        hdr_cells[0].text = f"{samples[i]['id']}({i})"
        hdr_cells[1].text = file
        hdr_cells[2].text = samples[i]["user_query"]
        hdr_cells[4].text = samples[i][f"{model}_response"]
        hdr_cells[6].text = "\n".join(
            [
                str(id) + ". " + fact
                for id, fact in enumerate(samples[i][f"{model}_fact"], 1)
            ]
        )
        para = document.add_paragraph()

        para.add_run("\n")
    print("saving to: " + save_path)
    document.save(save_path)


if __name__ == "__main__":
    file_list = [
        "Bio-Medical",
        "Finance",
        "Science",
        "Education",
        "Open-Domain",
    ]
    model = "llama-2-7b-chat-hf"
    json_dir = "./json/"
    save_dir = "./docs/"
    check_exist(save_dir)
    for file in file_list:
        data = read_json(path=os.path.join(json_dir, f"{file}.json"), part=200)
        data_lst = split_list(data, x=4)
        for i in range(len(data_lst)):
            save_path = os.path.join(save_dir, f"{file}_{i+1}.docx")
            generate_doc(model, file, data_lst[i], save_path)
